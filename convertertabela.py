import pandas as pd
from docx import Document

#MUDAR VALORES AQUI
INICIO = 552
FIM = 1233


def lerarquivoxlxs(nomearquivo):
    # Leitura do arquivo .xlsx
    df = pd.read_excel(nomearquivo, engine='openpyxl')  
    return df

def formatObjectID(objectid, format="P0000"):
    objectid = str(objectid)
    tamanho = len(format) - len(objectid)
    formatado = format[:tamanho] + objectid
    return formatado

def formatDirection(direction):
    #retornar xxº xx' xx''
    graus, minutos, segundos = direction.split("-")
    formatado = f"{graus}º {minutos}' {segundos}"
    formatado +='"'
    return formatado

def formatDistance(distance):
    return distance

def formatCoords(coords):
    formatado = f"{coords:,.4f}"
    formatado = formatado.replace(",", "X").replace(".", ",").replace("X", ".")
    formatado += " m"
    return formatado


df = lerarquivoxlxs("Vertices.xlsx")
print(df.head(18))
#print(df.tail())

inicio = 19
fim = 135
print("TESTE PADRÕES")
for n in range(inicio,fim):
    print(formatObjectID(df.loc[n, "OBJECTID_1"]))
    print(formatDirection(df.loc[n, "Direction"]))
    print(formatDistance(df.loc[n, "Distance"]))
    print(formatCoords(df.loc[n, "N"]))
    print(formatCoords(df.loc[n, "E"]))



inicio = INICIO-1
fim = FIM-1
df = lerarquivoxlxs("Vertices.xlsx")
print(df.head(fim))
# Cria o documento Word
doc = Document()
doc.add_heading("Sede Urbana "+formatObjectID(df.loc[inicio, "OBJECTID_1"])+" até "+formatObjectID(df.loc[fim, "OBJECTID_1"]), level=1)
paragrafo = doc.add_paragraph()


textoinicio = "Inicia-se a descrição deste perímetro no vértice "
paragrafo.add_run(textoinicio)
#Primeiro vértice
texto = formatObjectID(df.loc[inicio, "OBJECTID_1"]) + ", " + "de coordenadas N "+formatCoords(df.loc[inicio, "N"]) +" e E "+formatCoords(df.loc[inicio, "E"])+", que segue confrontando por linha seca em um azimute de "+ formatDirection(df.loc[inicio, "Direction"]) + " a uma distância de "+formatDistance(df.loc[inicio, "Distance"])+" até o vértice "
paragrafo.add_run(texto)

for n in range(inicio+1,fim+1):
    texto = formatObjectID(df.loc[n, "OBJECTID_1"]) + "; "
    paragrafo.add_run(texto)
    texto = "Do vértice " + formatObjectID(df.loc[n, "OBJECTID_1"]) + ", " + "de coordenadas N "+formatCoords(df.loc[n, "N"]) +" e E "+formatCoords(df.loc[n, "E"])+", que segue confrontando por linha seca em um azimute de "+ formatDirection(df.loc[n, "Direction"]) + " a uma distância de "+formatDistance(df.loc[n, "Distance"])+" até o vértice "
    paragrafo.add_run(texto)

#Adiciona último texto
texto =  " retornando ao vértice " + formatObjectID(df.loc[inicio, "OBJECTID_1"]) +", onde teve início essa descrição."
paragrafo.add_run(texto)

#Salva o documento
doc.save("saida-"+formatObjectID(df.loc[inicio, "OBJECTID_1"])+"-"+formatObjectID(df.loc[fim, "OBJECTID_1"])+".docx")

#produção de texto